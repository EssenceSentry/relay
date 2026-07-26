"""Render a source-cited Relay dossier from Markdown to DOCX and PDF.

The agent remains responsible for researching and writing the dossier. This
module deliberately performs no LLM calls: it validates the editorial
contract, preserves the supplied text and citations, and applies Blend360's
document design consistently.
"""

from __future__ import annotations

import hashlib
import re
import shutil
import subprocess
import tempfile
from collections.abc import Callable, Sequence
from dataclasses import dataclass
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from knowledge_core.ids import safe_filename

_SECTION_NAMES = (
    "Executive Summary",
    "The Challenge",
    "Our Solution",
    "Key Features",
    "Quantified Outcomes",
    "Business Value",
    "Relevant Visual Evidence",
    "Known Gaps / Caveats",
    "Sources Used",
)
_REQUIRED_SECTIONS = frozenset(_SECTION_NAMES) - {"Relevant Visual Evidence"}
_H1 = re.compile(r"^#\s+(?P<title>[^#].+?)\s*$")
_H2 = re.compile(r"^##\s+(?P<title>.+?)\s*$")
_BULLET = re.compile(r"^\s*[-*]\s+(?P<text>.+?)\s*$")
_LABELED_BULLET = re.compile(
    r"^\*\*(?P<label>.+?)\*\*(?:\s*[\u2014\u2013:-]\s*|\s+)(?P<text>.*)$"
)
_INLINE_MARKUP = re.compile(r"(\*\*.+?\*\*|`.+?`|\*[^*].*?\*)")
_GENERATED_LINE = re.compile(
    r"^Generated\s+(?P<date>.+?)(?:\s+[\u2014\u2013-]\s+.+)?$",
    re.IGNORECASE,
)
_METRIC_TOKEN = re.compile(r"(?:\d|[%$€£])")
_METRIC_LABEL = re.compile(
    r"^(?P<value>[~+<>=]?(?:[$€£])?(?:\d[\d,]*(?:\.\d+)?|\.\d+)"
    r"(?:[%A-Za-z]+)?)"
    r"(?:(?:\s+|-)(?P<label>.+))?$"
)
_MAX_MARKDOWN_CHARACTERS = 200_000
_MAX_TITLE_CHARACTERS = 300
_MAX_DESCRIPTOR_CHARACTERS = 500

NAVY = "062F53"
CYAN = "22B3E3"
INK = "29333C"
MUTED = "66737E"
PALE_CYAN = "EAF7FB"
PALE_AMBER = "FFF6E7"
WHITE = "FFFFFF"
BODY_FONT = "Arial"
DISPLAY_FONT = "Arial"


class DossierValidationError(ValueError):
    """The supplied Markdown does not satisfy the dossier contract."""


class DossierCompilationError(RuntimeError):
    """LaTeX could not produce a dossier PDF."""


@dataclass(frozen=True, slots=True)
class DossierBullet:
    label: str | None
    text: str

    @property
    def combined_text(self) -> str:
        if self.label and self.text:
            return f"{self.label} — {self.text}"
        return self.label or self.text


@dataclass(frozen=True, slots=True)
class DossierSection:
    paragraphs: tuple[str, ...]
    bullets: tuple[DossierBullet, ...]

    @property
    def is_empty(self) -> bool:
        return not self.paragraphs and not self.bullets


@dataclass(frozen=True, slots=True)
class DossierContent:
    title: str
    descriptor: str
    generated_date: str
    kicker: str
    sections: dict[str, DossierSection]

    def section(self, name: str) -> DossierSection:
        return self.sections[name]


@dataclass(frozen=True, slots=True)
class RenderedDossierFiles:
    markdown_path: Path
    latex_path: Path
    docx_path: Path
    pdf_path: Path
    source_sha256: str


@dataclass(frozen=True, slots=True)
class RenderedDossierBytes:
    markdown: bytes
    latex: bytes
    docx: bytes
    pdf: bytes
    source_sha256: str
    filename_stem: str


PdfCompiler = Callable[[Path], Path]


class LocalDossierRenderer:
    """Synchronous renderer used by the API service."""

    def render(
        self,
        markdown: str,
        *,
        filename_stem: str | None = None,
    ) -> RenderedDossierBytes:
        return render_dossier_bytes(
            markdown,
            filename_stem=filename_stem,
        )


def parse_dossier_markdown(markdown: str) -> DossierContent:
    """Validate and parse the Markdown contract used by the Relay skill."""

    normalized = markdown.replace("\r\n", "\n").replace("\r", "\n").strip()
    if not normalized:
        raise DossierValidationError("Dossier Markdown cannot be empty")
    if len(normalized) > _MAX_MARKDOWN_CHARACTERS:
        raise DossierValidationError(
            "Dossier Markdown exceeds the 200,000 character limit"
        )

    lines = normalized.splitlines()
    title_indexes = [
        index for index, line in enumerate(lines) if _H1.match(line)
    ]
    if title_indexes != [0]:
        raise DossierValidationError(
            "Dossier Markdown must begin with exactly one level-one title"
        )
    title_match = _H1.match(lines[0])
    assert title_match is not None
    title_line = title_match.group("title").strip()
    title, descriptor = _split_title(title_line)
    if len(title) > _MAX_TITLE_CHARACTERS:
        raise DossierValidationError(
            "Dossier title exceeds the 300 character limit"
        )
    if len(descriptor) > _MAX_DESCRIPTOR_CHARACTERS:
        raise DossierValidationError(
            "Dossier descriptor exceeds the 500 character limit"
        )

    preamble: list[str] = []
    section_lines: dict[str, list[str]] = {}
    current_section: str | None = None
    for line in lines[1:]:
        heading = _H2.match(line)
        if heading is not None:
            section_name = heading.group("title").strip()
            if section_name not in _SECTION_NAMES:
                expected = ", ".join(_SECTION_NAMES)
                raise DossierValidationError(
                    f"Unsupported dossier section {section_name!r}. "
                    f"Expected one of: {expected}"
                )
            if section_name in section_lines:
                raise DossierValidationError(
                    f"Dossier section {section_name!r} appears more than once"
                )
            current_section = section_name
            section_lines[current_section] = []
            continue
        if current_section is None:
            preamble.append(line)
        else:
            section_lines[current_section].append(line)

    missing = sorted(_REQUIRED_SECTIONS - section_lines.keys())
    if missing:
        raise DossierValidationError(
            "Dossier is missing required sections: " + ", ".join(missing)
        )

    sections = {
        name: _parse_section(section_lines.get(name, []))
        for name in _SECTION_NAMES
        if name in section_lines
    }
    empty_required = sorted(
        name for name in _REQUIRED_SECTIONS if sections[name].is_empty
    )
    if empty_required:
        raise DossierValidationError(
            "Required dossier sections cannot be empty: "
            + ", ".join(empty_required)
        )
    feature_count = len(sections["Key Features"].bullets)
    if not 1 <= feature_count <= 8:
        raise DossierValidationError(
            "Key Features must contain from 1 to 8 Markdown bullets"
        )
    if not sections["Sources Used"].bullets:
        raise DossierValidationError(
            "Sources Used must contain at least one Markdown bullet"
        )

    generated_date, kicker = _parse_preamble(preamble)
    return DossierContent(
        title=title,
        descriptor=descriptor,
        generated_date=generated_date,
        kicker=kicker,
        sections=sections,
    )


def render_dossier_files(
    markdown: str,
    output_directory: str | Path,
    *,
    filename_stem: str | None = None,
    pdf_compiler: PdfCompiler | None = None,
) -> RenderedDossierFiles:
    """Create Markdown, TeX, DOCX, and PDF files in one directory."""

    normalized = markdown.replace("\r\n", "\n").replace("\r", "\n").strip()
    content = parse_dossier_markdown(normalized)
    output = Path(output_directory).expanduser().resolve()
    output.mkdir(parents=True, exist_ok=True)
    stem = _filename_stem(filename_stem or content.title)
    markdown_path = output / f"{stem}.md"
    latex_path = output / f"{stem}.tex"
    docx_path = output / f"{stem}.docx"
    style_path = output / "blenddossier.sty"

    markdown_path.write_text(normalized + "\n", encoding="utf-8")
    style_source = (
        Path(__file__).resolve().parent
        / "assets"
        / "dossier"
        / "blenddossier.sty"
    )
    shutil.copy2(style_source, style_path)
    latex_path.write_text(
        render_latex_source(content),
        encoding="utf-8",
    )
    build_docx(content, docx_path)
    compiler = pdf_compiler or compile_pdf
    pdf_path = compiler(latex_path)
    if not pdf_path.is_file() or not pdf_path.read_bytes().startswith(b"%PDF-"):
        raise DossierCompilationError(
            "The LaTeX compiler did not produce a valid PDF"
        )
    return RenderedDossierFiles(
        markdown_path=markdown_path,
        latex_path=latex_path,
        docx_path=docx_path,
        pdf_path=pdf_path,
        source_sha256=hashlib.sha256(normalized.encode("utf-8")).hexdigest(),
    )


def render_dossier_bytes(
    markdown: str,
    *,
    filename_stem: str | None = None,
    pdf_compiler: PdfCompiler | None = None,
) -> RenderedDossierBytes:
    """Render a dossier into upload-ready byte strings."""

    with tempfile.TemporaryDirectory(prefix="relay-dossier-") as temp:
        files = render_dossier_files(
            markdown,
            temp,
            filename_stem=filename_stem,
            pdf_compiler=pdf_compiler,
        )
        return RenderedDossierBytes(
            markdown=files.markdown_path.read_bytes(),
            latex=files.latex_path.read_bytes(),
            docx=files.docx_path.read_bytes(),
            pdf=files.pdf_path.read_bytes(),
            source_sha256=files.source_sha256,
            filename_stem=files.markdown_path.stem,
        )


def compile_pdf(
    latex_path: Path,
    *,
    engine: str = "xelatex",
    runs: int = 2,
) -> Path:
    """Compile a generated dossier with the configured XeLaTeX runtime."""

    engine_path = shutil.which(engine)
    if engine_path is None:
        raise DossierCompilationError(
            f"{engine} is required to render dossier PDFs"
        )
    command = [
        engine_path,
        "-interaction=nonstopmode",
        "-halt-on-error",
        latex_path.name,
    ]
    for _ in range(runs):
        completed = subprocess.run(
            command,
            cwd=latex_path.parent,
            check=False,
            capture_output=True,
            text=True,
        )
        if completed.returncode != 0:
            tail = "\n".join(completed.stdout.splitlines()[-80:])
            raise DossierCompilationError(
                f"LaTeX compilation failed for {latex_path.name}:\n{tail}"
            )
    return latex_path.with_suffix(".pdf")


def build_docx(content: DossierContent, output_path: Path) -> Path:
    """Build an editable Word dossier using public python-docx APIs."""

    document = Document()
    _configure_document(document, content)
    _add_docx_title(document, content)
    _add_docx_page_one(document, content)
    document.add_page_break()
    _add_docx_page_two(document, content)
    document.add_page_break()
    _add_docx_page_three(document, content)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    return output_path


def render_latex_source(content: DossierContent) -> str:
    """Create the XeLaTeX source used for the dossier PDF."""

    summary = _latex_paragraphs(content.section("Executive Summary"))
    challenge = _latex_paragraphs(content.section("The Challenge"))
    solution = _latex_paragraphs(content.section("Our Solution"))
    features = "\n".join(
        "\\feature"
        f"{{{index:02d}}}"
        f"{{{latex_escape(item.label or f'Capability {index}')}}}"
        f"{{{latex_escape(item.text)}}}"
        for index, item in enumerate(
            content.section("Key Features").bullets,
            start=1,
        )
    )
    outcome_section = content.section("Quantified Outcomes")
    metrics, outcome_details = _split_metrics(outcome_section.bullets)
    metric_block = ""
    if metrics:
        metric_rows = "\n".join(
            f"\\metric{{{latex_escape(_metric_parts(item)[0])}}}"
            f"{{{latex_escape(_metric_parts(item)[1])}}}"
            for item in metrics
        )
        metric_block = (
            f"\\begin{{metricrow}}{{{len(metrics)}}}\n"
            f"{metric_rows}\n"
            "\\end{metricrow}\n"
        )
    outcome_body = _latex_paragraphs(
        DossierSection(
            outcome_section.paragraphs,
            metrics + outcome_details,
        )
    )
    business = _latex_paragraphs(content.section("Business Value"))
    visual = content.sections.get("Relevant Visual Evidence")
    visual_block = ""
    if visual is not None and not visual.is_empty:
        cards = "\n".join(
            "\\evidencecard"
            f"{{{latex_escape(item.label or 'Visual evidence')}}}"
            f"{{{latex_escape(item.text)}}}"
            f"{{{latex_escape(_citation_hint(item.combined_text))}}}"
            for item in _section_items(visual)
        )
        visual_block = (
            "\\dsection{Relevant Visual Evidence}\n"
            "\\begin{evidencerow}\n"
            f"{cards}\n"
            "\\end{evidencerow}\n\n"
        )
    caveats = "\n".join(
        f"\\caveat{{{index:02d}}}{{{latex_escape(item.combined_text)}}}"
        for index, item in enumerate(
            _section_items(content.section("Known Gaps / Caveats")),
            start=1,
        )
    )
    sources = "\n".join(
        f"\\source{{{index}}}{{{latex_escape(item.combined_text)}}}"
        for index, item in enumerate(
            content.section("Sources Used").bullets,
            start=1,
        )
    )
    source_columns = (
        f"\\begin{{multicols}}{{2}}\n{sources}\n\\end{{multicols}}"
        if len(content.section("Sources Used").bullets) > 3
        else sources
    )
    footer_date = latex_escape(content.generated_date.upper())
    return f"""\\documentclass[10pt]{{article}}
\\usepackage{{blenddossier}}

\\begin{{document}}
\\thispagestyle{{dossierfirst}}

\\dossierheader
  {{DOSSIER}}
  {{{latex_escape(content.title)}}}
  {{{latex_escape(content.descriptor)}}}

\\dossierfooter
  {{BLEND360 \\ \\textperiodcentered\\ PROJECT DOSSIER \\textperiodcentered\\ GENERATED {footer_date}}}
  {{CONFIDENTIAL \\textperiodcentered\\ INTERNAL SALES USE}}

\\dsection{{Executive Summary}}
{summary}

\\columnratio{{0.515}}
\\setlength{{\\columnsep}}{{24pt}}
\\begin{{paracol}}{{2}}

\\dsection{{The Challenge}}
{challenge}

\\dsection{{Our Solution}}
{solution}

\\switchcolumn

\\dsection{{Key Features}}
{features}

\\end{{paracol}}

\\newpage

\\dsection{{Quantified Outcomes}}
{metric_block}{outcome_body}

\\dsection{{Business Value}}
{business}

\\newpage

{visual_block}\\dsection{{Known Gaps / Caveats}}
\\begin{{caveatbox}}{{Known Gaps / Caveats}}
{caveats}
\\end{{caveatbox}}

\\dsection{{Sources Used}}
{source_columns}

\\end{{document}}
"""


def latex_escape(text: str) -> str:
    replacements = {
        "\\": r"\textbackslash{}",
        "&": r"\&",
        "%": r"\%",
        "$": r"\$",
        "#": r"\#",
        "_": r"\_",
        "{": r"\{",
        "}": r"\}",
        "~": r"\textasciitilde{}",
        "^": r"\textasciicircum{}",
    }
    plain = _plain_markdown(text)
    return "".join(
        replacements.get(character, character) for character in plain
    )


def _parse_section(lines: Sequence[str]) -> DossierSection:
    paragraphs: list[str] = []
    bullets: list[DossierBullet] = []
    paragraph_lines: list[str] = []

    def flush_paragraph() -> None:
        if paragraph_lines:
            paragraphs.append(" ".join(paragraph_lines).strip())
            paragraph_lines.clear()

    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            flush_paragraph()
            continue
        bullet = _BULLET.match(raw_line)
        if bullet is not None:
            flush_paragraph()
            bullets.append(_parse_bullet(bullet.group("text")))
            continue
        if line.startswith("### "):
            flush_paragraph()
            paragraph_lines.append(f"**{line[4:].strip()}**")
            continue
        paragraph_lines.append(line)
    flush_paragraph()
    return DossierSection(tuple(paragraphs), tuple(bullets))


def _parse_bullet(text: str) -> DossierBullet:
    labeled = _LABELED_BULLET.match(text.strip())
    if labeled is None:
        return DossierBullet(label=None, text=text.strip())
    return DossierBullet(
        label=labeled.group("label").strip(),
        text=labeled.group("text").strip(),
    )


def _split_title(title_line: str) -> tuple[str, str]:
    for separator in (" \N{EM DASH} ", " \N{EN DASH} ", " - "):
        if separator in title_line:
            title, descriptor = title_line.split(separator, maxsplit=1)
            if title.strip() and descriptor.strip():
                return title.strip(), descriptor.strip()
    return title_line.strip(), "Blend360 project dossier"


def _parse_preamble(lines: Sequence[str]) -> tuple[str, str]:
    generated = date.today().strftime("%B %d, %Y")
    kicker = "A Blend360 Case Study"
    for raw_line in lines:
        line = _plain_markdown(raw_line.strip())
        if not line:
            continue
        generated_match = _GENERATED_LINE.match(line)
        if generated_match is not None:
            generated = generated_match.group("date").strip()
        else:
            kicker = line
    return generated, kicker


def _filename_stem(value: str) -> str:
    cleaned = safe_filename(value).removesuffix(".md")
    return cleaned[:96] or "project-dossier"


def _plain_markdown(text: str) -> str:
    return text.replace("**", "").replace("`", "").replace("*", "").strip()


def _section_items(section: DossierSection) -> tuple[DossierBullet, ...]:
    paragraph_items = tuple(
        DossierBullet(label=None, text=paragraph)
        for paragraph in section.paragraphs
    )
    return paragraph_items + section.bullets


def _split_metrics(
    bullets: Sequence[DossierBullet],
) -> tuple[tuple[DossierBullet, ...], tuple[DossierBullet, ...]]:
    metrics: list[DossierBullet] = []
    details: list[DossierBullet] = []
    for bullet in bullets:
        label = bullet.label or ""
        if (
            len(metrics) < 4
            and label
            and len(label) <= 24
            and _METRIC_TOKEN.search(label)
        ):
            metrics.append(bullet)
        else:
            details.append(bullet)
    return tuple(metrics), tuple(details)


def _metric_parts(metric: DossierBullet) -> tuple[str, str]:
    label = metric.label or ""
    match = _METRIC_LABEL.fullmatch(label)
    if match is None:
        return label, "reported outcome"
    value = match.group("value")
    short_label = (match.group("label") or "reported outcome").strip()
    return value, short_label


def _citation_hint(text: str) -> str:
    match = re.search(r"\(([^()]+)\)\s*$", _plain_markdown(text))
    return match.group(1) if match is not None else "Source cited in dossier"


def _set_run_style(
    run: Run,
    *,
    size: float,
    color: str = INK,
    bold: bool = False,
    italic: bool = False,
    font_name: str = BODY_FONT,
) -> None:
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def _add_inline_markdown(
    paragraph: Paragraph,
    text: str,
    *,
    size: float,
    color: str = INK,
) -> None:
    position = 0
    for token in _INLINE_MARKUP.finditer(text):
        if token.start() > position:
            run = paragraph.add_run(text[position : token.start()])
            _set_run_style(run, size=size, color=color)
        value = token.group(0)
        bold = value.startswith("**")
        code = value.startswith("`")
        content = value[2:-2] if bold else value[1:-1]
        run = paragraph.add_run(content)
        _set_run_style(
            run,
            size=size,
            color=NAVY if bold else color,
            bold=bold or code,
            italic=not bold and not code,
        )
        position = token.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        _set_run_style(run, size=size, color=color)


def _configure_document(
    document: DocumentObject,
    content: DossierContent,
) -> None:
    document.core_properties.title = content.title
    document.core_properties.subject = "Blend360 project dossier"
    document.core_properties.author = "Blend360"
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.68)
    section.top_margin = Inches(0.64)
    section.bottom_margin = Inches(0.68)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_run_style(
        header.add_run("BLEND360  /  PROJECT DOSSIER"),
        size=7.5,
        color=MUTED,
        bold=True,
    )
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_run_style(
        footer.add_run(
            f"BLEND360  ·  GENERATED {content.generated_date.upper()}"
        ),
        size=7,
        color=MUTED,
        bold=True,
    )
    footer.add_run(" " * 8)
    _set_run_style(
        footer.add_run("CONFIDENTIAL · INTERNAL SALES USE"),
        size=7,
        color=CYAN,
        bold=True,
    )


def _add_docx_title(
    document: DocumentObject,
    content: DossierContent,
) -> None:
    tag = document.add_paragraph()
    tag.paragraph_format.space_after = Pt(7)
    _set_run_style(
        tag.add_run("PROJECT DOSSIER"),
        size=8,
        color=CYAN,
        bold=True,
    )
    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(5)
    title.paragraph_format.keep_with_next = True
    _set_run_style(
        title.add_run(content.title),
        size=24,
        color=NAVY,
        bold=True,
        font_name=DISPLAY_FONT,
    )
    descriptor = document.add_paragraph()
    descriptor.paragraph_format.space_after = Pt(11)
    descriptor.paragraph_format.keep_with_next = True
    _set_run_style(
        descriptor.add_run(content.descriptor.upper()),
        size=8.5,
        color=CYAN,
        bold=True,
    )


def _add_docx_page_one(
    document: DocumentObject,
    content: DossierContent,
) -> None:
    _add_docx_section(document, "Executive Summary")
    _add_docx_section_body(
        document,
        content.section("Executive Summary"),
        body_size=9.5,
    )
    columns = document.add_table(rows=1, cols=2)
    columns.alignment = WD_TABLE_ALIGNMENT.CENTER
    columns.autofit = False
    left, right = columns.rows[0].cells
    left.width = Inches(3.55)
    right.width = Inches(3.45)
    for cell in (left, right):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        cell.paragraphs[0].clear()

    _add_docx_section(left, "The Challenge")
    _add_docx_section_body(left, content.section("The Challenge"))
    _add_docx_section(left, "Our Solution")
    _add_docx_section_body(left, content.section("Our Solution"))

    _add_docx_section(right, "Key Features")
    for index, item in enumerate(
        content.section("Key Features").bullets,
        start=1,
    ):
        paragraph = right.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(6)
        _set_run_style(
            paragraph.add_run(f"{index:02d}  "),
            size=8,
            color=CYAN,
            bold=True,
        )
        if item.label:
            _set_run_style(
                paragraph.add_run(item.label),
                size=8.5,
                color=NAVY,
                bold=True,
            )
            if item.text:
                _set_run_style(
                    paragraph.add_run(f" — {item.text}"),
                    size=8.5,
                )
        else:
            _add_inline_markdown(
                paragraph,
                item.text,
                size=8.5,
            )


def _add_docx_page_two(
    document: DocumentObject,
    content: DossierContent,
) -> None:
    _add_docx_continuation(document, content)
    _add_docx_section(document, "Quantified Outcomes")
    outcomes = content.section("Quantified Outcomes")
    _add_docx_paragraphs(document, outcomes.paragraphs)
    metrics, details = _split_metrics(outcomes.bullets)
    if metrics:
        table = document.add_table(rows=1, cols=len(metrics))
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for cell, metric in zip(table.rows[0].cells, metrics, strict=True):
            metric_value, metric_label = _metric_parts(metric)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            value = cell.paragraphs[0]
            value.alignment = WD_ALIGN_PARAGRAPH.CENTER
            value.paragraph_format.space_after = Pt(4)
            _set_run_style(
                value.add_run(metric_value),
                size=19,
                color=NAVY,
                bold=True,
                font_name=DISPLAY_FONT,
            )
            detail = cell.add_paragraph()
            detail.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_inline_markdown(
                detail,
                metric_label,
                size=7.5,
                color=MUTED,
            )
        for metric in metrics:
            _add_docx_bullet(document, metric.combined_text)
    for item in details:
        _add_docx_bullet(document, item.combined_text)

    _add_docx_section(document, "Business Value")
    _add_docx_section_body(
        document,
        content.section("Business Value"),
        body_size=9.5,
    )


def _add_docx_page_three(
    document: DocumentObject,
    content: DossierContent,
) -> None:
    _add_docx_continuation(document, content)
    visual = content.sections.get("Relevant Visual Evidence")
    if visual is not None and not visual.is_empty:
        _add_docx_section(document, "Relevant Visual Evidence")
        items = _section_items(visual)
        table = document.add_table(
            rows=(len(items) + 2) // 3,
            cols=3,
        )
        table.style = "Light Shading Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for index, cell in enumerate(
            cell for row in table.rows for cell in row.cells
        ):
            if index >= len(items):
                cell.text = ""
                continue
            item = items[index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            heading = cell.paragraphs[0]
            heading.paragraph_format.space_after = Pt(4)
            _set_run_style(
                heading.add_run(item.label or "Visual evidence"),
                size=8.5,
                color=NAVY,
                bold=True,
            )
            body = cell.add_paragraph()
            _add_inline_markdown(body, item.text, size=7.5)

    _add_docx_section(document, "Known Gaps / Caveats")
    caveat_table = document.add_table(rows=1, cols=1)
    caveat_table.style = "Light Shading Accent 2"
    caveat_cell = caveat_table.cell(0, 0)
    caveat_cell.paragraphs[0].clear()
    for index, item in enumerate(
        _section_items(content.section("Known Gaps / Caveats")),
        start=1,
    ):
        paragraph = caveat_cell.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(5)
        _set_run_style(
            paragraph.add_run(f"{index:02d}  "),
            size=8,
            color="B96C16",
            bold=True,
        )
        _add_inline_markdown(
            paragraph,
            item.combined_text,
            size=8,
        )

    _add_docx_section(document, "Sources Used")
    for index, item in enumerate(
        content.section("Sources Used").bullets,
        start=1,
    ):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(4)
        _set_run_style(
            paragraph.add_run(f"{index}.  "),
            size=7.5,
            color=CYAN,
            bold=True,
        )
        _add_inline_markdown(
            paragraph,
            item.combined_text,
            size=7.5,
            color=MUTED,
        )


def _add_docx_continuation(
    document: DocumentObject,
    content: DossierContent,
) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_after = Pt(9)
    _set_run_style(
        paragraph.add_run(f"{content.title}  /  CONTINUED"),
        size=7.5,
        color=MUTED,
        bold=True,
    )


def _add_docx_section(container: Any, title: str) -> None:
    paragraph = container.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(7)
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.keep_with_next = True
    _set_run_style(
        paragraph.add_run("■  "),
        size=7,
        color=CYAN,
        bold=True,
    )
    _set_run_style(
        paragraph.add_run(title.upper()),
        size=8.5,
        color=NAVY,
        bold=True,
    )


def _add_docx_section_body(
    container: Any,
    section: DossierSection,
    *,
    body_size: float = 8.5,
) -> None:
    _add_docx_paragraphs(
        container,
        section.paragraphs,
        body_size=body_size,
    )
    for item in section.bullets:
        _add_docx_bullet(
            container,
            item.combined_text,
            body_size=body_size,
        )


def _add_docx_paragraphs(
    container: Any,
    paragraphs: Sequence[str],
    *,
    body_size: float = 8.5,
) -> None:
    for text in paragraphs:
        paragraph = container.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(5)
        paragraph.paragraph_format.line_spacing = 1.12
        _add_inline_markdown(paragraph, text, size=body_size)


def _add_docx_bullet(
    container: Any,
    text: str,
    *,
    body_size: float = 8.5,
) -> None:
    paragraph = container.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.12)
    paragraph.paragraph_format.first_line_indent = Inches(-0.12)
    paragraph.paragraph_format.space_after = Pt(4)
    _set_run_style(
        paragraph.add_run("—  "),
        size=body_size,
        color=CYAN,
        bold=True,
    )
    _add_inline_markdown(paragraph, text, size=body_size)


def _latex_paragraphs(section: DossierSection) -> str:
    chunks = [
        f"\\body{{{latex_escape(paragraph)}}}"
        for paragraph in section.paragraphs
    ]
    chunks.extend(
        f"\\outcome{{{latex_escape(item.combined_text)}}}"
        for item in section.bullets
    )
    return "\n".join(chunks)
